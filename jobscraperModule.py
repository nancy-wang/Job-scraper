
#A web scraper to do job searches for Health Leads clients 
#This code was adapted from https://github.com/Futvin/Web-Scraper-in-Python-with-BeautifulSoup/blob/master/runSimpleScrap.py
import re
import urllib.request
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt


document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Garamond'
font.size = Pt(12)
#new_heading_style = document.styles['Heading 1']
#font_heading = new_heading_style.font
#font_heading.name = 'Garamond'
#font_heading.size = Pt(24)
document.add_heading('Job Listings', 1)

def crawler_func(url, page, jobsubj):
	job_subj = jobsubj.replace("+", "")
	with urllib.request.urlopen(url) as urldata:
		rawtext = urldata.read().decode('utf-8', 'ignore')
		soup = BeautifulSoup(rawtext, 'html.parser')


	jobListings = soup.select(".result")
	for job in jobListings:
		#Get job details
		jobTitle= job.select(".jobtitle")[0].text.lstrip().rstrip()
		companyName = job.select(".company")[0].text.lstrip().rstrip()
		jobLocation = job.select(".location")[0].text.lstrip().rstrip()
		
		# Figured out if this ad is sponsored.
		isSponsered = ""
		sponderedSpans = job.find_all("span")
		if len(sponderedSpans) > 0:
			for span in sponderedSpans:
				spanText = span.text.lstrip().rstrip()
				if re.match("Sponsored", spanText) is not None:
					isSponsered = spanText

   #get job URL
		all_jobs = job.find_all("a")
		for link in all_jobs:
   			try:
   				if "turnstileLink" in link['class']:
   					jobURL = "https://www.indeed.com" + link['href']

   			except:
   				pass

		try:
			if (page < 10):
   				if(isSponsered != "Sponsored"):
   					p = document.add_paragraph()
   					p.add_run(jobTitle).bold = True
   					q = document.add_paragraph(companyName)
   					q.add_run(" " + jobLocation)
   					document.add_paragraph(jobURL)
   					document.add_paragraph()
   			
		except:
   			pass
	document.save('job_listings_' + str(job_subj) + '.docx')
	
	if (page < 10):
		document.add_paragraph("If applicable, other jobs you might want to consider are linked below.")
	relatedJobs = soup.select(".relatedQuerySpacing")
	for related in relatedJobs:
		all_related = related.find_all("a")
		
		for link in all_related:
			try:
				relatedURL = "https://www.indeed.com" + link['href']
				document.add_paragraph(relatedURL)
			except:
				pass
	document.save('job_listings_' + str(job_subj) + '.docx')
	try:
		next_link = soup.select(".pagination")[0].find_all('a')

		for link in next_link:
   			if re.match("Next", link.text) is not None:
   				page += 10
   				print(link['href'])
   				crawler_func("https://www.indeed.com" + link['href'],page, jobsubj)
	except:
		pass

def menu():
	print("Hey fellow Employment Specialist! Welcome to the job scraper!")
	zipcode = input("Please input the zipcode of the job you are looking for.\nIf you don't know what zipcode you're looking for, enter 'N/A' and the default location will be set to Baltimore\n")
	if (zipcode == "N/A"):
		zipcode = 21218
	jobsubj = input("Please input the type of job you're seeking.\nIf your input is more than one word, please separate spaces with + symbols, e.g. 'night+jobs'\n")
	radius = input("Please input how far you can travel.\n If no preference, input 25.\n")
	sortdate = input("Type 'y' if you'd like to sort by date or 'n' if not. Otherwise jobs will be sorted by relevance.\n")
	if (sortdate=='y'):
		url = "https://www.indeed.com/jobs?q=" + str(jobsubj) + "&l=" + str(zipcode) + "&radius=" + str(radius) + "&sort=date"
	else: 
		url = "https://www.indeed.com/jobs?q=" + str(jobsubj) + "&l=" + str(zipcode) + "&radius=" + str(radius)
	#count = 1
	print("Thanks! Starting search now...")
	print(url)
	page = 0
	crawler_func(url, page, jobsubj)
	print("Jobs written to Word document. Thanks for using this service!")

menu()



